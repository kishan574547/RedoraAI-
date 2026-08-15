import io
import os
import base64
from typing import Dict, Any, Optional
import docx
from PIL import Image
import pytesseract
from app.services.pdf_tools import extract_text as extract_pdf_text
from app.services.openrouter_client import OpenRouterClient
from app.core.logging import logger

# Configure pytesseract path if installed on Windows default paths
possible_tesseract_paths = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    r"C:\Users\Public\Tesseract-OCR\tesseract.exe"
]
for p in possible_tesseract_paths:
    if os.path.exists(p):
        pytesseract.pytesseract.tesseract_cmd = p
        break


class DocumentScannerService:
    def __init__(self):
        self.openrouter_client = OpenRouterClient()

    async def _describe_image_with_vision(self, file_bytes: bytes, filename: str, content_type: str) -> str:
        """Fallback vision call using OpenRouter vision model to describe image content."""
        try:
            b64_img = base64.b64encode(file_bytes).decode("utf-8")
            media_type = content_type if content_type else "image/png"
            if not media_type.startswith("image/"):
                media_type = "image/png"
            
            data_url = f"data:{media_type};base64,{b64_img}"

            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"Please analyze and provide a detailed summary of this image ('{filename}'). Describe all visible objects, text, charts, diagrams, or key information present in the image."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": data_url
                            }
                        }
                    ]
                }
            ]

            # Try free vision capable models on OpenRouter
            vision_models = [
                "google/gemini-2.5-flash",
                "qwen/qwen-2.5-vl-72b-instruct",
                "meta-llama/llama-3.2-11b-vision-instruct"
            ]

            for model_name in vision_models:
                try:
                    res = await self.openrouter_client.chat_completion(messages=messages, model=model_name, max_tokens=600)
                    content = res["choices"][0]["message"]["content"]
                    if content and len(content.strip()) > 10:
                        logger.info(f"Image vision description generated successfully using {model_name}")
                        return f"[Visual Image Analysis of '{filename}']:\n{content.strip()}"
                except Exception as ve:
                    logger.warning(f"Vision model {model_name} failed: {str(ve)}")
                    continue

            return f"[Image '{filename}' attached. OCR detected minimal text.]"
        except Exception as e:
            logger.error(f"Failed to generate vision description: {str(e)}")
            return f"[Image '{filename}' attached.]"

    async def extract_content(self, file_bytes: bytes, filename: str, content_type: str = "") -> Dict[str, Any]:
        """
        Extract text from file based on type:
        - PDF: pdf_tools extract_text
        - DOCX: python-docx
        - TXT: utf-8 text
        - Image (JPG/PNG): pytesseract OCR -> vision model description fallback if OCR yields little/no text
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            try:
                pdf_res = extract_pdf_text(file_bytes)
                extracted_text = pdf_res.get("full_text", "").strip()
                return {
                    "file_type": "pdf",
                    "extracted_text": extracted_text or "[Empty PDF Document]",
                    "status": "success"
                }
            except Exception as e:
                logger.error(f"Error parsing PDF '{filename}': {str(e)}")
                raise ValueError(f"Failed to parse PDF file: {str(e)}")

        elif ext in [".docx", ".doc"]:
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                extracted_text = "\n\n".join(full_text).strip()
                return {
                    "file_type": "docx",
                    "extracted_text": extracted_text or "[Empty Word Document]",
                    "status": "success"
                }
            except Exception as e:
                logger.error(f"Error parsing DOCX '{filename}': {str(e)}")
                raise ValueError(f"Failed to parse Word document: {str(e)}")

        elif ext in [".txt"]:
            try:
                extracted_text = file_bytes.decode("utf-8", errors="ignore").strip()
                return {
                    "file_type": "txt",
                    "extracted_text": extracted_text or "[Empty Text File]",
                    "status": "success"
                }
            except Exception as e:
                logger.error(f"Error reading TXT '{filename}': {str(e)}")
                raise ValueError(f"Failed to read text file: {str(e)}")

        elif ext in [".jpg", ".jpeg", ".png"]:
            try:
                img = Image.open(io.BytesIO(file_bytes))
                ocr_text = ""
                try:
                    ocr_text = pytesseract.image_to_string(img).strip()
                except Exception as ocr_err:
                    logger.warning(f"Tesseract OCR engine unavailable or failed: {str(ocr_err)}")

                # Check if OCR extracted meaningful text (at least 20 chars)
                if ocr_text and len(ocr_text) >= 20:
                    logger.info(f"OCR extracted {len(ocr_text)} chars from '{filename}'")
                    extracted_text = f"[OCR Extracted Text from Image '{filename}']:\n{ocr_text}"
                else:
                    logger.info(f"OCR extracted minimal text ({len(ocr_text)} chars). Using vision fallback for '{filename}'")
                    vision_desc = await self._describe_image_with_vision(file_bytes, filename, content_type)
                    if ocr_text:
                        extracted_text = f"[OCR Extracted Partial Text]:\n{ocr_text}\n\n{vision_desc}"
                    else:
                        extracted_text = vision_desc

                return {
                    "file_type": "image",
                    "extracted_text": extracted_text,
                    "status": "success"
                }
            except Exception as e:
                logger.error(f"Error processing image '{filename}': {str(e)}")
                raise ValueError(f"Failed to process image file: {str(e)}")

        else:
            raise ValueError(f"Unsupported file format '{ext}'. Accepted formats are PDF, DOCX, TXT, JPG, PNG.")


document_scanner = DocumentScannerService()
