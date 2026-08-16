import io
import re
from typing import List, Dict, Any, Optional
import docx
import pypdf

from app.services.openrouter_client import OpenRouterClient
from app.core.logging import logger


ACTION_VERBS = {
    "developed", "created", "built", "implemented", "engineered", "designed",
    "led", "managed", "spearheaded", "architected", "optimized", "increased",
    "reduced", "achieved", "delivered", "automated", "launched", "executed",
    "collaborated", "orchestrated", "transformed", "streamlined", "configured",
    "established", "formulated", "headed", "initiated", "pioneered", "resolved"
}

STOP_WORDS = {
    "a", "an", "the", "and", "or", "but", "if", "then", "else", "when", "at",
    "from", "by", "for", "with", "about", "against", "between", "into", "through",
    "during", "before", "after", "above", "below", "to", "in", "on", "off",
    "over", "under", "again", "further", "that", "this", "these", "those", "is",
    "are", "was", "were", "be", "been", "being", "have", "has", "had", "do",
    "does", "did", "doing", "would", "should", "could", "ought", "im", "youre",
    "his", "her", "their", "our", "your", "my", "we", "you", "they", "them",
    "will", "can", "shall", "must", "work", "job", "candidate", "role", "position"
}


async def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text content from PDF, DOCX, or Image file bytes using multi-stage extractors + Google Gemini Vision OCR."""
    fname = filename.lower()
    text = ""

    if fname.endswith(".pdf"):
        # Stage 1: pypdf extraction
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                t = page.extract_text() or ""
                pages_text.append(t)
            text = "\n".join(pages_text).strip()
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {str(e)}")

        # Stage 2: Fallback to pdf2docx if text is empty or very short (< 30 chars)
        if not text or len(text) < 30:
            try:
                import tempfile
                import os
                from pdf2docx import Converter
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                    f.write(file_bytes)
                    temp_pdf_path = f.name
                temp_docx_path = temp_pdf_path + ".docx"
                try:
                    cv = Converter(temp_pdf_path)
                    cv.parse(temp_docx_path)
                    cv.close()
                    doc = docx.Document(temp_docx_path)
                    docx_paras = [p.text for p in doc.paragraphs if p.text.strip()]
                    for table in doc.tables:
                        for row in table.rows:
                            row_t = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                            if row_t:
                                docx_paras.append(row_t)
                    text = "\n".join(docx_paras).strip()
                finally:
                    if os.path.exists(temp_pdf_path):
                        try: os.remove(temp_pdf_path)
                        except Exception: pass
                    if os.path.exists(temp_docx_path):
                        try: os.remove(temp_docx_path)
                        except Exception: pass
            except Exception as e:
                logger.warning(f"pdf2docx fallback extraction failed: {str(e)}")

        # Stage 3: Vision OCR fallback for scanned/image PDFs (e.g. Canva exports or photo PDFs)
        if not text or len(text) < 30:
            try:
                import fitz  # PyMuPDF
                import base64
                from app.services.openrouter_client import OpenRouterClient

                doc = fitz.open(stream=file_bytes, filetype="pdf")
                vision_pages = []
                client = OpenRouterClient()

                for page_num in range(min(len(doc), 5)):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(dpi=150)
                    img_png = pix.tobytes("png")
                    b64_img = base64.b64encode(img_png).decode("utf-8")

                    messages = [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": "Transcribe all text from this resume page image word for word. Extract every section heading, contact info (email, phone, location), experience, education, skills, projects, and bullet points."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{b64_img}"
                                    }
                                }
                            ]
                        }
                    ]

                    res = await client.chat_completion(messages=messages, model="gemini-flash-lite-latest", max_tokens=1500)
                    if isinstance(res, dict) and res.get("choices"):
                        p_text = res["choices"][0]["message"]["content"]
                        if p_text and len(p_text.strip()) > 15:
                            vision_pages.append(p_text.strip())

                if vision_pages:
                    text = "\n\n".join(vision_pages).strip()
                    logger.info(f"Successfully extracted {len(text)} characters from PDF '{filename}' via Google Gemini Vision OCR")
            except Exception as ve:
                logger.warning(f"Google Gemini Vision OCR failed for PDF '{filename}': {str(ve)}")

    elif fname.endswith(".docx") or fname.endswith(".doc"):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            text = "\n".join(paragraphs).strip()
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise ValueError("Failed to extract text from Word document. File may be corrupted.")
    elif fname.endswith((".png", ".jpg", ".jpeg")):
        try:
            import base64
            from app.services.openrouter_client import OpenRouterClient
            b64_img = base64.b64encode(file_bytes).decode("utf-8")
            media_type = "image/png" if fname.endswith(".png") else "image/jpeg"
            client = OpenRouterClient()
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Transcribe all text from this resume image word for word. Extract every section heading, contact info (email, phone, location), experience, education, skills, projects, and bullet points."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{b64_img}"
                            }
                        }
                    ]
                }
            ]
            res = await client.chat_completion(messages=messages, model="gemini-flash-lite-latest", max_tokens=2000)
            if isinstance(res, dict) and res.get("choices"):
                text = res["choices"][0]["message"]["content"].strip()
        except Exception as ie:
            logger.error(f"Error reading Image file: {str(ie)}")
            raise ValueError(f"Failed to extract text from image resume: {str(ie)}")
    else:
        raise ValueError("Unsupported file format. Please upload a PDF (.pdf), Word document (.docx), or Image (.jpg/.png).")

    return text.strip()


def rule_based_checks(text: str) -> Dict[str, Any]:
    """Perform rule-based ATS compliance checks on resume text."""
    lower_text = text.lower()
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    words = re.findall(r"\b\w+\b", lower_text)
    word_count = len(words)

    # 0. Text Extractability Check
    text_length = len(text.strip())
    is_text_readable = text_length >= 100

    # 1. Essential Sections Check
    sections = {
        "Experience": bool(re.search(r"\b(experience|work history|employment|history)\b", lower_text)),
        "Education": bool(re.search(r"\b(education|academic|qualification|university|college)\b", lower_text)),
        "Skills": bool(re.search(r"\b(skills|technical skills|competencies|technologies)\b", lower_text)),
        "Projects / Certifications": bool(re.search(r"\b(projects|certifications|awards|accomplishments)\b", lower_text))
    }
    missing_sections = [sec for sec, present in sections.items() if not present]

    # 2. Contact Info Check
    has_email = bool(re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text))
    has_phone = bool(re.search(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text))

    # 3. Bullet Point Usage
    bullet_lines = [line for line in lines if line.startswith(("•", "-", "*", "▪", "➢", "v")) or re.match(r"^\d+\.", line)]
    bullet_count = len(bullet_lines)

    # 4. Action Verbs Check
    first_words = set()
    for b in bullet_lines:
        b_clean = re.sub(r"^[•\-\*\▪\➢v\d\.\s]+", "", b).strip()
        parts = b_clean.split()
        if parts:
            first_words.add(parts[0].lower())
    matched_action_verbs = list(first_words.intersection(ACTION_VERBS))

    # 5. Quantifiable Metrics Check
    metrics = re.findall(r"\b\d+(?:[\.,]\d+)?%|\$\d+(?:,\d+)*(?:\.\d+)?|\b\d+\+\b|\b\d+k\b", lower_text)

    # 6. Word Count Check
    optimal_length = 300 <= word_count <= 1200

    items = [
        {
            "name": "ATS Text Extractability",
            "passed": is_text_readable,
            "details": f"Extracted {text_length} characters ({word_count} words). Text is clear and readable by ATS parsers." if is_text_readable else f"Extracted 0 readable characters. Your PDF appears to be a scanned image or non-text PDF. Real ATS systems (Workday, Greenhouse) cannot parse image PDFs and will score your resume 0%. Convert to a text-selectable PDF or Word document."
        },
        {
            "name": "Essential Sections Present",
            "passed": len(missing_sections) == 0,
            "details": f"Found sections. Missing: {', '.join(missing_sections)}" if missing_sections else "All key sections (Experience, Education, Skills, Projects) found."
        },
        {
            "name": "Contact Details Detectable",
            "passed": has_email and has_phone,
            "details": f"Email: {'Found' if has_email else 'Missing'}, Phone: {'Found' if has_phone else 'Missing'}."
        },
        {
            "name": "Bullet Points Usage",
            "passed": bullet_count >= 4,
            "details": f"Detected {bullet_count} bullet points for clean formatting."
        },
        {
            "name": "Action Verbs Starting Bullets",
            "passed": len(matched_action_verbs) >= 3,
            "details": f"Used {len(matched_action_verbs)} strong action verbs ({', '.join(matched_action_verbs[:5])})."
        },
        {
            "name": "Quantifiable Metrics & Results",
            "passed": len(metrics) >= 2,
            "details": f"Found {len(metrics)} quantifiable metrics/percentages (e.g. {', '.join(metrics[:3])})." if metrics else "No clear metrics or percentages found. Add numbers to show impact."
        },
        {
            "name": "Optimal Word Count Length",
            "passed": optimal_length,
            "details": f"Resume is {word_count} words ({'Optimal 300-1200 range' if optimal_length else 'Too short or too long'})."
        }
    ]

    passed_count = sum(1 for item in items if item["passed"])
    section_score = round((passed_count / len(items)) * 100, 1)

    return {
        "score": section_score,
        "word_count": word_count,
        "items": items
    }


def keyword_match_score(resume_text: str, job_description: str) -> Dict[str, Any]:
    """Calculate keyword overlap between resume text and job description."""
    if not job_description or not job_description.strip():
        return {"match_percentage": None, "missing_keywords": []}

    jd_words = re.findall(r"\b[a-zA-Z]{3,}\b", job_description.lower())
    res_words = set(re.findall(r"\b[a-zA-Z]{3,}\b", resume_text.lower()))

    # Frequency analysis for job description keywords
    freq: Dict[str, int] = {}
    for w in jd_words:
        if w not in STOP_WORDS:
            freq[w] = freq.get(w, 0) + 1

    # Sort top 25 keywords by frequency
    sorted_keywords = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:25]
    top_jd_keywords = [k for k, v in sorted_keywords]

    matched_keywords = [k for k in top_jd_keywords if k in res_words]
    missing_keywords = [k for k in top_jd_keywords if k not in res_words]

    match_pct = round((len(matched_keywords) / max(1, len(top_jd_keywords))) * 100, 1)

    return {
        "match_percentage": match_pct,
        "matched_keywords": matched_keywords[:10],
        "missing_keywords": missing_keywords[:12]
    }


async def ai_qualitative_feedback(resume_text: str, job_description: str = "") -> List[str]:
    """Generate qualitative ATS feedback using OpenRouter LLM, with intro preamble stripping."""
    if len(resume_text.strip()) < 30:
        return [
            "Convert Image/Scanned PDF to Text PDF: Your uploaded document contains 0 readable text characters. Re-export your resume directly from Word, Google Docs, or Canva as a text-selectable PDF.",
            "Avoid Graphic Layouts: Do not export resumes as PNG/JPG images before saving to PDF, as ATS systems cannot parse image text.",
            "Include Standard Headings: Use standard section titles (Experience, Education, Skills, Projects) so parsers can index your content.",
            "Quantify Achievements: Use numbers, percentages, and metrics to demonstrate concrete accomplishments."
        ]

    client = OpenRouterClient()
    
    prompt = (
        "You are an expert ATS (Applicant Tracking System) Resume Specialist. "
        "Analyze the following resume text and provide exactly 4 concise, actionable bullet points for improvement. "
        "Do NOT include introductory phrases like 'Certainly!', 'Below are', 'Here are', or numbering headers. "
        "Start every bullet directly with an action verb or recommendation topic.\n\n"
        f"Resume Text:\n{resume_text[:2500]}\n"
    )

    if job_description:
        prompt += f"\nTarget Job Description:\n{job_description[:1500]}\n"

    try:
        res = await client.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=500
        )
        content = ""
        if isinstance(res, dict):
            choices = res.get("choices", [])
            if choices and isinstance(choices[0], dict):
                content = choices[0].get("message", {}).get("content", "")

        if content:
            preamble_triggers = ("certainly", "below are", "here are", "improve formatting", "sure", "here is", "based on", "recommendations:")
            raw_lines = content.split("\n")
            cleaned_lines = []
            for line in raw_lines:
                clean_l = line.strip(" *-•0123456789.:")
                if not clean_l or len(clean_l) < 15:
                    continue
                if any(clean_l.lower().startswith(p) for p in preamble_triggers):
                    continue
                cleaned_lines.append(clean_l)

            if cleaned_lines:
                return cleaned_lines[:4]
    except Exception as e:
        logger.error(f"Error fetching AI qualitative feedback: {str(e)}")

    # Intelligent fallback feedback if LLM is offline
    fallback = [
        "Include measurable impact (numbers, percentages, revenue, time saved) in every experience bullet point.",
        "Ensure standard ATS section headers (Experience, Education, Skills, Projects) are used without icons or tables.",
        "Begin all bullet points under work experience with strong past-tense action verbs (e.g., Engineered, Spearheaded, Optimized).",
        "Tailor technical skills directly to target job requirements to improve keyword index matching."
    ]
    return fallback


async def generate_custom_ai_suggestions(
    resume_text: str,
    custom_instruction: str,
    job_description: str = ""
) -> List[str]:
    """Generate custom ATS suggestions based on user prompt/instruction using OpenRouter LLM."""
    if not custom_instruction.strip():
        return []

    client = OpenRouterClient()
    
    prompt = (
        "You are an expert ATS Resume Coach. "
        f"The candidate has provided the following custom request/instruction: '{custom_instruction.strip()}'.\n\n"
        "Analyze the candidate's resume text below and generate 2-3 specific, highly targeted ATS optimization suggestions fulfilling their request.\n"
        "Do NOT include introductory filler (e.g. 'Certainly!', 'Here are'). Start each bullet directly with an action verb or topic.\n\n"
        f"Resume Text:\n{resume_text[:2500]}\n"
    )

    if job_description:
        prompt += f"\nTarget Job Description:\n{job_description[:1500]}\n"

    try:
        res = await client.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=400
        )
        content = ""
        if isinstance(res, dict):
            choices = res.get("choices", [])
            if choices and isinstance(choices[0], dict):
                content = choices[0].get("message", {}).get("content", "")

        if content:
            preamble_triggers = ("certainly", "below are", "here are", "sure", "here is", "based on", "recommendations:")
            raw_lines = content.split("\n")
            cleaned_lines = []
            for line in raw_lines:
                clean_l = line.strip(" *-•0123456789.:")
                if not clean_l or len(clean_l) < 12:
                    continue
                if any(clean_l.lower().startswith(p) for p in preamble_triggers):
                    continue
                cleaned_lines.append(clean_l)

            if cleaned_lines:
                return cleaned_lines[:3]
    except Exception as e:
        logger.error(f"Error generating custom AI suggestions: {str(e)}")

    # Fallback suggestion based on custom instruction
    return [
        f"Incorporate user request: '{custom_instruction.strip()}' throughout your Experience and Skills sections.",
        "Refine work achievements to highlight leadership and technical metrics aligned with your target goals."
    ]


async def apply_resume_corrections(
    original_text: str,
    accepted_suggestions: List[str],
    missing_keywords: List[str]
) -> str:
    """Rewrite and optimize resume text based on user-accepted suggestions and missing keywords."""
    client = OpenRouterClient()
    
    prompt = (
        "You are an expert ATS Resume Editor. "
        "Rewrite the following resume text into a clean, highly optimized ATS resume format.\n\n"
        "Instructions:\n"
        "1. Incorporate the following user-accepted suggestions:\n" + "\n".join(f"- {s}" for s in accepted_suggestions) + "\n\n"
        "2. Seamlessly weave in the following missing keywords where relevant:\n" + ", ".join(missing_keywords) + "\n\n"
        "3. Use standard ATS section headings: EXPERIENCE, EDUCATION, SKILLS, PROJECTS.\n"
        "4. Format work achievements into action-verb bullet points with quantifiable impact.\n"
        "5. Output ONLY the polished resume text without any conversational preamble or markdown codeblock wrappers.\n\n"
        f"Original Resume Text:\n{original_text[:3000]}"
    )

    try:
        res = await client.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1500
        )
        content = ""
        if isinstance(res, dict):
            choices = res.get("choices", [])
            if choices and isinstance(choices[0], dict):
                content = choices[0].get("message", {}).get("content", "")

        if content and len(content.strip()) > 100:
            return content.strip()
    except Exception as e:
        logger.error(f"Failed to apply LLM resume corrections: {str(e)}")

    # Fallback rewriter if LLM is offline
    lines = [l.strip() for l in original_text.split("\n") if l.strip()]
    header = lines[0] if lines else "CANDIDATE RESUME"
    
    enhanced = [
        header.upper(),
        "=" * len(header),
        "",
        "SKILLS",
        "------",
        f"Core Competencies: {', '.join(missing_keywords[:8]) if missing_keywords else 'Technical Leadership, Project Management'}",
        "",
        "EXPERIENCE",
        "----------"
    ]

    for line in lines[1:]:
        if line.startswith(("•", "-", "*")):
            enhanced.append(line)
        else:
            enhanced.append(f"• {line}")

    if accepted_suggestions:
        enhanced.append("")
        enhanced.append("ADDITIONAL OPTIMIZATIONS APPLIED:")
        for s in accepted_suggestions:
            enhanced.append(f"• Applied: {s}")

    return "\n".join(enhanced)


def generate_optimized_docx(resume_text: str) -> bytes:
    """Generate a clean ATS-friendly Word (.docx) document from text."""
    doc = docx.Document()

    # Set standard 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)

    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    if not lines:
        lines = ["OPTIMIZED RESUME", "No content provided."]

    # Header / Title
    header_p = doc.add_paragraph()
    run = header_p.add_run(lines[0])
    run.font.name = "Arial"
    run.font.size = docx.shared.Pt(18)
    run.bold = True
    header_p.paragraph_format.space_after = docx.shared.Pt(12)

    for line in lines[1:]:
        if line.isupper() and len(line) < 30:
            h = doc.add_heading(level=2)
            h_run = h.add_run(line)
            h_run.font.name = "Arial"
            h_run.font.size = docx.shared.Pt(13)
            h_run.bold = True
            h.paragraph_format.space_before = docx.shared.Pt(12)
            h.paragraph_format.space_after = docx.shared.Pt(4)
        elif line.startswith(("•", "-", "*", "▪", "➢")):
            clean_b = re.sub(r"^[•\-\*\▪\➢\s]+", "", line).strip()
            p = doc.add_paragraph(style='List Bullet')
            p_run = p.add_run(clean_b)
            p_run.font.name = "Arial"
            p_run.font.size = docx.shared.Pt(10.5)
            p.paragraph_format.space_after = docx.shared.Pt(3)
        else:
            p = doc.add_paragraph()
            p_run = p.add_run(line)
            p_run.font.name = "Arial"
            p_run.font.size = docx.shared.Pt(10.5)
            p.paragraph_format.space_after = docx.shared.Pt(4)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_optimized_pdf(resume_text: str) -> bytes:
    """Generate a clean ATS-friendly PDF document from text using ReportLab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=12
    )

    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#0369a1'),
        spaceBefore=10,
        spaceAfter=4
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=4
    )

    bullet_style = ParagraphStyle(
        'BulletCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )

    story = []
    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    if not lines:
        lines = ["OPTIMIZED RESUME", "No content provided."]

    # Header / Name
    story.append(Paragraph(lines[0].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), title_style))

    for line in lines[1:]:
        clean_text = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        if line.isupper() and len(line) < 30:
            story.append(Paragraph(clean_text, heading_style))
        elif line.startswith(("•", "-", "*", "▪", "➢")):
            b_clean = re.sub(r"^[•\-\*\▪\➢\s]+", "", clean_text).strip()
            story.append(Paragraph(f"• {b_clean}", bullet_style))
        else:
            story.append(Paragraph(clean_text, body_style))

    doc.build(story)
    return buffer.getvalue()


